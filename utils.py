"""
Description: 
    
-*- Encoding: UTF-8 -*-
@File     ：utils.py.py
@Author   ：King Songtao
@Time     ：2024/12/16 下午3:59
@Contact  ：king.songtao@gmail.com
"""
import traceback
from datetime import date
from docx.shared import Pt
import re
import streamlit as st
import sqlite3
import pandas as pd
import time

from sqlalchemy import text


def format_date(input_date):
    """
    自定义日期格式化
    将 2024-12-11 格式化为 11th Dec. 2024
    """
    day = input_date.day
    month_abbr = input_date.strftime('%b')
    year = input_date.year

    # 添加日期后缀
    if 10 <= day % 100 <= 20:
        suffix = 'th'
    else:
        suffix = {1: 'st', 2: 'nd', 3: 'rd'}.get(day % 10, 'th')

    return f"{day}{suffix} {month_abbr}. {year}"


def replace_placeholders(doc, data_dict):
    """
    替换文档中的占位符并统一字体
    """
    # 替换段落中的文本
    for paragraph in doc.paragraphs:
        for key, value in data_dict.items():
            if key in paragraph.text:
                paragraph.text = paragraph.text.replace(key, str(value))

        # 统一段落字体和大小
        for run in paragraph.runs:
            run.font.name = 'Arial'
            run.font.size = Pt(10)

    # 替换表格中的文本
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for key, value in data_dict.items():
                    if key in cell.text:
                        cell.text = cell.text.replace(key, str(value))

                # 统一单元格内段落的字体和大小
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.font.name = 'Arial'
                        run.font.size = Pt(10)

    return doc


def stream_res(res):
    """前端制作流式输出效果"""
    for char in res:
        yield char
        time.sleep(0.02)


def extract_date_from_html(html_content):
    """
    提取 HTML 内容中的日期
    假设日期格式为 "16th Dec. 2024"（可以根据需要调整正则表达式）
    """
    # 用正则表达式匹配日期
    date_pattern = r"\d{1,2}[a-zA-Z]{2}\s[A-Za-z]+\.\s\d{4}"
    match = re.search(date_pattern, html_content)

    if match:
        return match.group(0)  # 返回匹配到的日期
    return None


def validate_address(address):
    """
    验证地址是否只包含英文字符、数字和常见标点符号
    """
    # 增加 / 到允许的字符中
    pattern = r'^[a-zA-Z0-9\s\.\,\-\#\/]+$'

    if not address:
        return False, "地址不能为空"

    if not re.match(pattern, address):
        return False, "地址只能包含英文字符、数字和符号(.,- #/)。请检查是否含有中文逗号！"

    return True, ""


def connect_db():
    try:
        conn = st.connection('mysql', type='sql')
        print(f"数据库连接成功！")
        return conn
    except Exception as e:
        print(f"数据库连接失败，错误信息：{e}")
        return None


# def connect_db():
#     conn = sqlite3.connect(r"./work_orders.db")
#     return conn


def insert_data_to_db(register_time, notes, work_time, address, project, dispatcher, confirmed, registered, dispatched, dispatch_price, final_price, receipt_or_invoice):
    conn = connect_db()
    with conn.session as s:
        try:
            insert_query = text(
                """
                INSERT INTO work_orders (
                    record_time, notes, work_time, address, basic_plan, 
                    dispatcher, confirmed, registered, dispatched, 
                    sales_price, final_price, receipt
                )
                VALUES (:record_time, :notes, :work_time, :address, :basic_plan, :dispatcher, :confirmed, :registered, :dispatched, :sales_price, :final_price, :receipt)       
                """
            )

            params = {
                "record_time": register_time,
                "notes": notes,
                "work_time": work_time,
                "address": address,
                "basic_plan": project,
                "dispatcher": dispatcher,
                "confirmed": confirmed,
                "registered": registered,
                "dispatched": dispatched,
                "sales_price": dispatch_price,
                "final_price": final_price,
                "receipt": receipt_or_invoice
            }
            s.execute(insert_query, params)
            s.commit()
            success = st.success("工单创建成功！页面将在3秒后刷新...", icon="✅")
            time.sleep(1)
            success.empty()
            success = st.success("工单创建成功！页面将在2秒后刷新...", icon="✅")
            time.sleep(1)
            success.empty()
            success = st.success("工单创建成功！页面将在1秒后刷新...", icon="✅")
            time.sleep(1)
            st.cache_data.clear()
            st.rerun()
        except Exception as e:
            st.error(f"添加失败，错误原因：{e}")


# 显示工单数据
def display_all_orders():
    conn = connect_db()
    df = conn.query("SELECT * FROM work_orders", ttl=600)
    st.dataframe(df, hide_index=False)


def display_preview_data():
    conn = connect_db()
    df = conn.query("SELECT * FROM work_orders", ttl=600)
    df = df[['work_time', 'dispatcher', 'address', 'final_price', 'sales_price', 'notes']]
    from config import CUSTOM_HEADER
    df = df.rename(columns=CUSTOM_HEADER)
    st.dataframe(df, hide_index=True, use_container_width=True, on_select="ignore", selection_mode="single-row")


def edit_work_order_page(record_time, notes, work_time, address, basic_plan, dispatcher, confirmed, registered, dispatched, sales_price, final_price, receipt):
    conn = connect_db()
    with conn.session as s:
        try:
            update_query = text(
                "UPDATE work_orders SET record_time = :record_time , notes = :notes, work_time = :work_time, address = :address, basic_plan = :basic_plan, dispatcher = :dispatcher, confirmed = :confirmed, registered = :registered, dispatched = :dispatched, sales_price = :sales_price, final_price = :final_price, receipt = :receipt WHERE address = :address")
            param = {
                "record_time": record_time,
                "notes": notes,
                "work_time": work_time,
                "address": address,
                "basic_plan": basic_plan,
                "dispatcher": dispatcher,
                "confirmed": confirmed,
                "registered": registered,
                "dispatched": dispatched,
                "sales_price": sales_price,
                "final_price": final_price,
                "receipt": receipt,
            }
            s.execute(update_query, params=param)
            s.commit()
            success = st.success("工单更新成功！页面将在3秒后刷新...", icon="✅")
            time.sleep(1)
            success.empty()
            success = st.success("工单更新成功！页面将在2秒后刷新...", icon="✅")
            time.sleep(1)
            success.empty()
            success = st.success("工单更新成功！页面将在1秒后刷新...", icon="✅")
            time.sleep(1)
            st.cache_data.clear()
            st.rerun()
        except Exception as e:
            st.error(f"更新失败，错误原因：{e}")


def format_work_orders(data):
    formatted_orders = []
    for i in range(len(data['id'])):
        formatted_order = f"{data['work_time'][i]} | {data['dispatcher'][i]} | {data['address'][i]}"
        formatted_orders.append(formatted_order)
    return formatted_orders


def get_all_addresses():
    # 从数据库获取所有工单地址
    # 示例：查询数据库返回所有地址

    conn = connect_db()  # 请替换成你的数据库文件路径
    addresses = conn.query("SELECT id, work_time, dispatcher, address  FROM work_orders", ttl=600)
    order_info = format_work_orders(addresses.to_dict())
    return order_info


def get_order_by_address(order_info):
    address = order_info.split(" | ")[2]
    conn = connect_db()  # 请替换成你的数据库文件路径
    result = conn.query("SELECT * FROM work_orders WHERE address = :address", params={'address': address}, ttl=360)

    return result.to_dict()


def delete_work_order(address):
    conn = connect_db()
    with conn.session as s:
        try:
            delete_query = text("DELETE FROM work_orders WHERE address = :address")

            s.execute(delete_query, params={'address': address})
            s.commit()
            success = st.success("工单删除成功！页面将在3秒后刷新...", icon="✅")
            time.sleep(1)
            success.empty()
            success = st.success("工单删除成功！页面将在2秒后刷新...", icon="✅")
            time.sleep(1)
            success.empty()
            success = st.success("工单删除成功！页面将在1秒后刷新...", icon="✅")
            time.sleep(1)
            st.cache_data.clear()
            st.rerun()
        except Exception as e:
            st.warning(f"失败, 错误信息：{e}")


def get_total_sale():
    conn = connect_db()
    try:
        query_query = """
        SELECT SUM(final_price) AS total_final_price FROM work_orders 
        """
        result = conn.query(query_query, ttl=600)
        result = result['total_final_price'][0]
        return result
    except Exception as e:
        st.error(f"查询失败，错误原因：{e}")
        return None
